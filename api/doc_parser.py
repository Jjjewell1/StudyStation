"""Extract plain text from uploaded documents and split it into chunks.

Supported types: PDF, DOCX, TXT/MD, CSV, JSON, HTML. Anything else raises
UnsupportedDocError so the API can respond 4xx before storing anything.
Chunking is character-based on word boundaries with a small overlap so a
natural-language retriever (tsvector rank) can still find spans that straddle
a chunk boundary.
"""

from __future__ import annotations

import csv
import html.parser
import io
import json
import os

CHUNK_CHARS = 3000
CHUNK_OVERLAP = 150

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".html", ".htm"}


class UnsupportedDocError(ValueError):
    """The file type has no text extractor (e.g. .png, .pages, .epub)."""


class EmptyExtractionError(ValueError):
    """The file parsed but produced no text (e.g. a scanned-image PDF)."""


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    import pdfplumber

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages = []
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_html(data: bytes) -> str:
    class _TextHarvest(html.parser.HTMLParser):
        def __init__(self):
            super().__init__()
            self.parts: list[str] = []
            self.skip_depth = 0

        def handle_starttag(self, tag, attrs):
            if tag in ("script", "style"):
                self.skip_depth += 1

        def handle_endtag(self, tag):
            if tag in ("script", "style") and self.skip_depth:
                self.skip_depth -= 1

        def handle_data(self, data):
            if not self.skip_depth and data.strip():
                self.parts.append(data)

    parser = _TextHarvest()
    parser.feed(_decode(data))
    return "\n".join(parser.parts)


def _extract_csv(data: bytes) -> str:
    text = _decode(data)
    try:
        rows = list(csv.reader(io.StringIO(text)))
    except csv.Error:
        return text
    return "\n".join("\t".join(row) for row in rows)


def extract_text(filename: str, data: bytes) -> str:
    """Return extracted plain text for *filename*, or raise."""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocError(
            f"Unsupported file type '{ext or '(none)'}'. Supported: "
            + ", ".join(sorted(SUPPORTED_EXTENSIONS))
        )

    if ext == ".pdf":
        text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    elif ext == ".html" or ext == ".htm":
        text = _extract_html(data)
    elif ext == ".csv":
        text = _extract_csv(data)
    else:
        text = _decode(data)

    text = text.strip()
    if not text:
        raise EmptyExtractionError(
            f"'{filename}' produced no readable text (scanned/image-only PDF?)."
        )
    return text


def chunk_text(text: str) -> list[str]:
    """Split text into up-to-CHUNK_CHARS pieces on word boundaries."""
    if not text:
        return []
    words = text.split(" ")
    chunks: list[str] = []
    cur: list[str] = []
    cur_len = 0
    overlap_words: list[str] = []
    overlap_len = 0

    def _tail(words_: list[str], budget: int) -> tuple[list[str], int]:
        out: list[str] = []
        n = 0
        for w in reversed(words_):
            if n + len(w) + 1 > budget:
                break
            out.append(w)
            n += len(w) + 1
        return out[::-1], n

    for word in words:
        if cur_len + len(word) + 1 > CHUNK_CHARS and cur:
            chunks.append(" ".join(cur))
            overlap_words, overlap_len = _tail(cur, CHUNK_OVERLAP)
            cur, cur_len = list(overlap_words), overlap_len
        cur.append(word)
        cur_len += len(word) + 1
    if cur:
        chunks.append(" ".join(cur).strip())
    return [c.strip() for c in chunks if c.strip()]